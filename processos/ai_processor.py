"""
processos/ai_processor.py

Substituto direto do seu arquivo atual. Mantém os mesmos nomes públicos
(AIProcessor, extract_text_from_file, process_with_ai, process_directory,
merge_results), então nada mais no projeto precisa mudar.

O que muda por dentro:
  * PDF e imagem vão INTEIROS para a API (plugin file-parser da OpenRouter).
    Não há mais pré-extração de texto para esses formatos — é justamente a
    pré-extração que perdia tabela e devolvia vazio em PDF escaneado.
  * XLSX/DOCX/ODT/EML continuam sendo lidos localmente: nenhuma API de IA
    lê esses formatos: o parsing é responsabilidade sua.
  * O motor de PDF é escolhido por documento: `pdf-text` (grátis) quando há
    camada de texto, `native`/`mistral-ocr` quando é escaneado.
  * >>> NOVO: PDF escaneado passa antes pelo OCR local (ocrmypdf/Tesseract),
    que carimba a camada de texto de volta NO PDF. O arquivo continua indo
    inteiro para a API — o layout da tabela é preservado — mas pelo motor
    grátis. O motor pago só entra quando o OCR local falha ou rende pouco.

Dependências de sistema (Ubuntu 22.04, tudo via apt, sem pip):
    sudo apt install python3-xlrd ocrmypdf tesseract-ocr-por poppler-utils
"""
import base64
import email
import email.policy
import json
import logging
import mimetypes
import os
import random
import re
import shutil          # >>> NOVO
import subprocess      # >>> NOVO
import tempfile        # >>> NOVO
import time
from pathlib import Path

import requests
from django.conf import settings
from django.core.exceptions import ImproperlyConfigured

logger = logging.getLogger(__name__)

_SEM_PRECO_AINDA = {"", "*", "-", "--"}

# =====================================================================
# Compatibilidade de PDF entre três gerações da biblioteca
# ---------------------------------------------------------------------
# Ubuntu 22.04 (apt install python3-pypdf2) entrega PyPDF2 1.26.0, onde a
# classe é PdfFileReader e os métodos são getNumPages()/extractText().
# A classe PdfReader NÃO existe nessa versão — era exatamente esse o erro
# do código anterior: o import trazia PdfFileReader e o corpo usava PdfReader.
#
# Ordem de preferência: pypdf (moderno) > PyPDF2 2.x/3.x > PyPDF2 1.26.
# =====================================================================
def abrir_pdf(file_path):
    """Devolve (lista_de_textos_por_pagina, total_de_paginas)."""
    leitor = None
    api_legada = False

    try:                                   # pypdf 3.x / 5.x (pip)
        from pypdf import PdfReader
        leitor = PdfReader(file_path)
    except ImportError:
        try:                               # PyPDF2 2.x / 3.x (apt noble 24.04)
            from PyPDF2 import PdfReader
            leitor = PdfReader(file_path)
        except ImportError:                # PyPDF2 1.26 (apt jammy 22.04)
            from PyPDF2 import PdfFileReader
            leitor = PdfFileReader(file_path, strict=False)
            api_legada = True

    if api_legada:
        total = leitor.getNumPages()
        paginas = []
        for i in range(total):
            try:
                paginas.append(leitor.getPage(i).extractText() or "")
            except Exception:                              # noqa: BLE001
                logger.debug("Página %s ilegível em %s", i, file_path)
                paginas.append("")
        return paginas, total

    paginas = []
    for numero, pagina in enumerate(leitor.pages):
        try:
            paginas.append(pagina.extract_text() or "")
        except Exception:                                  # noqa: BLE001
            logger.debug("Página %s ilegível em %s", numero, file_path)
            paginas.append("")
    return paginas, len(paginas)


EXTENSOES_IMAGEM = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}


# =====================================================================
# >>> NOVO: OCR local (ocrmypdf + Tesseract), tudo via apt
# ---------------------------------------------------------------------
# Estratégia: NÃO extrair texto solto — texto solto perde a coluna, e
# proposta é tabela. O ocrmypdf devolve o MESMO PDF com uma camada de
# texto posicionada sobre a imagem, então o documento segue inteiro para
# a API e o motor cai no `pdf-text`, que é grátis.
#
# Devolve o caminho do PDF ocerizado (em diretório temporário, que o
# chamador apaga) ou None quando não deu — aí o fluxo antigo assume.
# =====================================================================
def ocr_pdf(file_path, idiomas=None, timeout=None):
    if not shutil.which("ocrmypdf"):
        return None

    pasta = tempfile.mkdtemp(prefix="ocr_")
    destino = os.path.join(pasta, os.path.basename(file_path))
    comando = [
        "ocrmypdf",
        "-l", idiomas or getattr(settings, "OCR_IDIOMAS", "por+eng"),
        "--skip-text",                    # página que já tem texto passa intacta
        "--rotate-pages", "--deskew",     # scanner torto / foto de proposta
        "--optimize", "0",                # não depende de jbig2enc/pngquant
        "--tesseract-timeout", str(getattr(settings, "OCR_TIMEOUT_PAGINA", 180)),
        "--quiet",
        file_path, destino,
    ]

    try:
        subprocess.run(comando, check=True,
                       timeout=timeout or getattr(settings, "OCR_TIMEOUT_TOTAL", 900),
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except (subprocess.SubprocessError, OSError):
        shutil.rmtree(pasta, ignore_errors=True)
        return None

    if not os.path.exists(destino):
        shutil.rmtree(pasta, ignore_errors=True)
        return None
    return destino


# =====================================================================
# >>> NOVO: leitura determinística do Modelo de Proposta (MB / COMRJ)
# ---------------------------------------------------------------------
# Serve para as duas pontas:
#   * modelo em branco enviado às empresas -> LINHA DE BASE (itens/PI)
#   * modelo devolvido preenchido          -> cotação, sem custo de IA
# Devolve o MESMO dicionário do merge_results, ou None se o layout não
# for reconhecido (aí o arquivo segue para a IA, como hoje).
# =====================================================================
_CAB_MODELO = {
    "item": "item", "numero de estoque": "pi", "nr de estoque": "pi",
    "nomenclatura": "descricao", "u.f.": "uf", "uf": "uf",
    "qt.": "qtde", "qt": "qtde", "qtde": "qtde",
    "prazo entrega": "prazo_entrega",
    "preco unit. (r$)": "preco_unitario", "preco unitario": "preco_unitario",
    "preco total (r$)": "preco_total", "preco total": "preco_total",
}
_ROTULOS_EMPRESA = {
    "razao social": "nome", "cnpj": "cnpj", "codemp": "codemp",
    "e-mail": "email", "email": "email", "telefone": "telefone",
}
_RUIDO_MODELO = {"", "*", "mrc", "marca", "local de entrega", "referencia",
                 "descricao caracteristica", "resposta decodificada",
                 "--------------------"}
_ACENTOS = str.maketrans("áàâãäéèêëíìîïóòôõöúùûüç", "aaaaaeeeeiiiiooooouuuuc")


def _norm_modelo(valor):
    return re.sub(r"\s+", " ", str(valor or "").strip().lower().translate(_ACENTOS))


def _ler_grade(file_path):
    """Matriz de células. Números inteiros viram string sem '.0' (PI!)."""
    if file_path.lower().endswith(".xls"):
        import xlrd
        aba = xlrd.open_workbook(file_path).sheet_by_index(0)

        def celula(r, c):
            cel = aba.cell(r, c)
            if cel.ctype == xlrd.XL_CELL_NUMBER and float(cel.value).is_integer():
                return str(int(cel.value))
            return cel.value

        return [[celula(r, c) for c in range(aba.ncols)] for r in range(aba.nrows)]

    from openpyxl import load_workbook
    wb = load_workbook(file_path, data_only=True)
    grade = [list(linha) for linha in wb.worksheets[0].iter_rows(values_only=True)]
    wb.close()
    return grade


def parse_modelo_proposta(file_path, limite_descricao=500):
    from .services import para_float          # import tardio: evita ciclo

    if Path(file_path).suffix.lower() not in (".xls", ".xlsx", ".xlsm"):
        return None
    try:
        grade = _ler_grade(file_path)
    except Exception:
        return None

    # --- acha a linha de cabeçalho da tabela de itens ------------------
    colunas, linha_cab = {}, None
    for i, linha in enumerate(grade[:80]):
        achadas = {}
        for c, valor in enumerate(linha):
            campo = _CAB_MODELO.get(_norm_modelo(valor))
            if campo and campo not in achadas:
                achadas[campo] = c
        if {"item", "pi", "descricao"} <= set(achadas):
            colunas, linha_cab = achadas, i
            break
    if linha_cab is None:
        return None

    def cel(linha, campo):
        c = colunas.get(campo)
        return "" if c is None or c >= len(linha) else str(linha[c] or "").strip()

    # --- identificação da empresa (só existe na proposta devolvida) ----
    identificacao = {}
    for linha in grade[:linha_cab]:
        for c, valor in enumerate(linha):
            campo = _ROTULOS_EMPRESA.get(_norm_modelo(valor))
            if not campo or identificacao.get(campo):
                continue
            for d in range(c + 1, min(c + 6, len(linha))):
                preenchido = str(linha[d] or "").strip()
                if preenchido and preenchido != "*":
                    identificacao[campo] = preenchido
                    break
    nome_empresa = identificacao.get("nome", "")

    # --- blocos de item: a linha do item começa com o número ----------
    itens, atual = [], None
    for linha in grade[linha_cab + 1:]:
        primeiro = cel(linha, "item")
        if re.fullmatch(r"\d{1,4}([.,]0+)?", primeiro):
            preco_txt = cel(linha, "preco_unitario")
            preco = None if _norm_modelo(preco_txt) in _SEM_PRECO_AINDA else para_float(preco_txt)
            atual = {
                "item": int(float(primeiro.replace(",", "."))),
                "pi": cel(linha, "pi"),
                "nsn": "", "codigo": cel(linha, "pi"),
                "nome_em_portugues": cel(linha, "descricao"),
                "qtde": para_float(cel(linha, "qtde")),
                "uf": cel(linha, "uf"),
                "prazo_entrega": cel(linha, "prazo_entrega"),
                "empresas": {nome_empresa: preco} if (nome_empresa and preco) else {},
                "valor_unitario_estimado": None,
                # na proposta devolvida o total é da EMPRESA, não da estimativa
                "valor_total": None if nome_empresa else para_float(cel(linha, "preco_total")),
                "confianca": 100,
                "avisos": [],
            }
            itens.append(atual)
            continue

        # linhas de continuação da nomenclatura / descrição característica
        if atual is None or len(atual["nome_em_portugues"]) >= limite_descricao:
            continue
        extra = cel(linha, "descricao")
        if extra and _norm_modelo(extra) not in _RUIDO_MODELO and not extra.startswith("<<"):
            atual["nome_em_portugues"] = (
                atual["nome_em_portugues"] + " " + extra).strip()[:limite_descricao]

    if not itens:
        return None

    avisos = []
    sem_pi = [i["item"] for i in itens if not i["pi"]]
    if sem_pi:
        avisos.append("itens sem PI no modelo (não entram no mapa): "
                      + ", ".join(map(str, sem_pi)))
    if nome_empresa and not any(i["empresas"] for i in itens):
        identificacao["tipo_resposta"] = "declinio"
        avisos.append(f"{nome_empresa} devolveu o modelo sem nenhum preço")

    return {
        "informacoes_gerais": {},
        "empresas": ([dict(identificacao,
                           tipo_resposta=identificacao.get("tipo_resposta", "cotacao"))]
                     if nome_empresa else []),
        "itens": itens,
        "avisos_gerais": avisos,
        "origem": "modelo-deterministico",
    }


# =====================================================================
# Extração do JSON da resposta
# =====================================================================
def extrair_json(texto):
    """
    Varredura com contagem de chaves, ignorando o que está dentro de aspas.
    O `re.search(r'\\{.*\\}')` da versão anterior quebrava quando havia
    chave dentro de uma string JSON (acontece em descrição de item).
    """
    if not texto:
        raise ValueError("resposta vazia")

    t = texto.strip()
    if t.startswith("```"):
        t = t.split("\n", 1)[-1]
        if t.rstrip().endswith("```"):
            t = t.rstrip()[:-3]
    t = t.strip()

    try:
        return json.loads(t)
    except json.JSONDecodeError:
        pass

    inicio = t.find("{")
    if inicio == -1:
        raise ValueError("nenhum objeto JSON na resposta")

    profundidade = 0
    em_string = False
    escape = False
    for i in range(inicio, len(t)):
        c = t[i]
        if escape:
            escape = False
            continue
        if c == "\\":
            escape = True
            continue
        if c == '"':
            em_string = not em_string
            continue
        if em_string:
            continue
        if c == "{":
            profundidade += 1
        elif c == "}":
            profundidade -= 1
            if profundidade == 0:
                return json.loads(t[inicio:i + 1])

    # Resposta truncada por max_tokens: tenta fechar as chaves abertas.
    try:
        return json.loads(t[inicio:].rstrip().rstrip(",") + "}" * profundidade)
    except json.JSONDecodeError as exc:
        raise ValueError(f"JSON truncado e irrecuperável: {exc}") from exc


# =====================================================================
class AIProcessor:

    SYSTEM_PROMPT = (
        "Você é um especialista em análise de documentos de licitação e compras "
        "públicas da Marinha do Brasil. Extraia informações estruturadas para "
        "preencher um Mapa Comparativo de Preços (Lei nº 14.133/2021, "
        "IN SEGES/ME nº 65/2021).\n"
        "REGRAS:\n"
        "1. NUNCA invente dados. Campo ausente no documento = string vazia.\n"
        "2. NUNCA calcule preço que não esteja escrito. Se só houver o total e "
        "a quantidade, preencha valor_total e deixe valor_unitario vazio.\n"
        "3. Números em formato brasileiro (1.234,56) viram float (1234.56).\n"
        "4. O PI (NÚMERO DE ESTOQUE) é a chave única do item. Copie-o "
        "exatamente como está, sem reformatar.\n"
        "5. Quando houver LISTA DE PI abaixo, use SOMENTE esses PI. Item do "
        "documento que não casar com nenhum: não invente linha, registre em "
        "avisos_gerais o que foi lido.\n"
        "6. Classifique cada empresa em tipo_resposta: 'cotacao' (há ao menos "
        "um preço), 'declinio' (recusa, não fornece, sem interesse, fora de "
        "linha, sem estoque) ou 'duvida' (só pergunta/pedido de "
        "esclarecimento). Um mesmo e-mail pode cotar parte e declinar o "
        "resto: nesse caso é 'cotacao' e o motivo vai em motivo_declinio.\n"
        "7. Perguntas do fornecedor vão em 'perguntas', nunca viram preço.\n"
        "8. Se a imagem/página estiver ilegível, registre em avisos.\n"
        # >>> NOVO: o documento pode ter vindo de OCR local
        "9. Se o texto vier de OCR e algum número estiver ambíguo (dígito "
        "borrado, separador decimal duvidoso, PI com caractere trocado), "
        "reduza 'confianca' do item e descreva a dúvida em avisos. Não "
        "'conserte' número por conta própria.\n"
        "10. Responda APENAS com o objeto JSON, sem markdown e sem preâmbulo."
    )

    ESQUEMA = """{
  "informacoes_gerais": {"numero_processo":"","modalidade":"","objeto":"",
                         "data_documento":"","valor_estimado_total":null},
  "empresas": [{"nome":"","cnpj":"","codemp":"","email":"","telefone":"",
                "tipo_resposta":"cotacao|declinio|duvida","motivo_declinio":"",
                "validade_proposta":"","prazo_entrega":"","valor_global":null}],
  "itens": [{
     "item": 1, "pi": "", "nsn": "", "codigo": "",
     "nome_em_portugues": "", "qtde": null, "uf": "",
     "empresas": {"NOME EXATO DA EMPRESA": 1234.56},
     "valor_unitario_estimado": null, "valor_total": null,
     "confianca": 90, "avisos": []
  }],
  "perguntas": [{"empresa":"","pergunta":""}],
  "avisos_gerais": []
}"""

    def __init__(self):
        # getattr com mensagem própria: sem isto, settings mal configurado
        # estoura AttributeError cru no meio do processamento de um pacote.
        self.api_key = getattr(settings, "OPENROUTER_API_KEY", "")
        if not self.api_key:
            raise ImproperlyConfigured(
                "OPENROUTER_API_KEY não configurada. Defina-a nas variáveis de"
                "ambiente do servidor — nunca no código versionado."
            )
        self.base_url = getattr(settings, "OPENROUTER_BASE_URL",
                                "https://openrouter.ai/api/v1")
        self.model = getattr(settings, "OPENROUTER_MODEL", "")
        # (conexao, leitura): falha rapido se a rede bloqueia, mas espera
        # o tempo necessario quando o PDF grande esta sendo processado.
        self.timeout = (
            getattr(settings, "OPENROUTER_CONNECT_TIMEOUT", 300),
            getattr(settings, "OPENROUTER_TIMEOUT", 3000),
        )
        # Rede corporativa/militar: proxy explicito e CA propria.
        # None faz o requests cair nas variaveis HTTP_PROXY/HTTPS_PROXY.
        self.proxies = getattr(settings, "OPENROUTER_PROXIES", None)
        self.verify = getattr(settings, "OPENROUTER_CA_BUNDLE", True)
        self.pdf_engine_scan = getattr(settings, "OPENROUTER_PDF_ENGINE_SCAN", "native")
        # >>> NOVO: OCR local antes de recorrer ao motor pago
        self.ocr_local = getattr(settings, "OCR_LOCAL", True)

    # -----------------------------------------------------------------
    # Roteamento da entrada
    # -----------------------------------------------------------------
    def montar_conteudo(self, file_path):
        """
        Decide como o arquivo entra na requisição.
        Devolve (blocos_de_conteudo, plugins, rota).
        """
        ext = Path(file_path).suffix.lower()

        # --- PDF: vai inteiro, o file-parser cuida ---------------------
        if ext == ".pdf":
            caminho, engine, ocerizado = file_path, "pdf-text", False

            if not self._pdf_tem_texto(file_path):
                # >>> NOVO: tenta o OCR local antes de pagar OCR/visão
                saida = ocr_pdf(file_path) if self.ocr_local else None
                if saida and self._pdf_tem_texto(saida):
                    caminho, ocerizado = saida, True
                else:
                    # OCR local ausente, falhou ou rendeu quase nada:
                    # fluxo antigo, motor pago sobre o arquivo original.
                    if saida:
                        shutil.rmtree(os.path.dirname(saida), ignore_errors=True)
                    engine = self.pdf_engine_scan
            try:
                blocos = [{
                    "type": "file",
                    "file": {
                        "filename": os.path.basename(file_path),
                        "file_data": "data:application/pdf;base64," + self._b64(caminho),
                    },
                }]
            finally:
                # O temporário do OCR sai do disco mesmo se a leitura falhar.
                if ocerizado:
                    shutil.rmtree(os.path.dirname(caminho), ignore_errors=True)

            plugins = [{"id": "file-parser", "pdf": {"engine": engine}}]
            return blocos, plugins, f"pdf/{engine}{'+ocr-local' if ocerizado else ''}"

        # --- imagem: direto para o modelo multimodal -------------------
        if ext in EXTENSOES_IMAGEM:
            mime = mimetypes.guess_type(file_path)[0] or "image/png"
            blocos = [{
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64," + self._b64(file_path)},
            }]
            return blocos, None, "imagem"

        # --- demais formatos: parsing local ----------------------------
        dados = self.extract_text_from_file(file_path)
        return ([{"type": "text", "text": dados["content"]}], None,
                f"texto/{ext.lstrip('.') or 'sem-extensao'}")

    @staticmethod
    def _b64(file_path):
        with open(file_path, "rb") as f:
            return base64.b64encode(f.read()).decode()

    # >>> NOVO: poppler primeiro. O extractText() da PyPDF2 1.26 subestima
    # muito o texto, e aqui isso custa caro: PDF digital classificado como
    # escaneado dispara OCR à toa (ou o motor pago). O pdftotext do
    # poppler-utils é o mesmo que o Ubuntu já instala com o Evince.
    @staticmethod
    def _texto_das_paginas(file_path, amostra=5):
        if shutil.which("pdftotext"):
            try:
                saida = subprocess.run(
                    ["pdftotext", "-f", "1", "-l", str(amostra), "-q", file_path, "-"],
                    check=True, timeout=120,
                    stdout=subprocess.PIPE, stderr=subprocess.DEVNULL)
                paginas = saida.stdout.decode("utf-8", "ignore").split("\f")
                if paginas and paginas[-1] == "":
                    paginas.pop()          # o \f final não é uma página
                if paginas:
                    return paginas[:amostra]
            except (subprocess.SubprocessError, OSError):
                pass

        paginas, _total = abrir_pdf(file_path)
        return paginas[:amostra]

    @classmethod
    def _pdf_tem_texto(cls, file_path, minimo_por_pagina=40, amostra=5):
        """
        Amostra as primeiras páginas. Com camada de texto usa o motor grátis;
        sem camada, tenta OCR local e, em último caso, paga OCR/visão.

        Usada duas vezes por PDF escaneado: antes do OCR (decide se precisa)
        e depois (decide se o OCR rendeu o bastante para valer o motor grátis).
        """
        try:
            paginas = cls._texto_das_paginas(file_path, amostra)
            if not paginas:
                return False
            soma = sum(len(t.strip()) for t in paginas)
            return (soma / len(paginas)) >= minimo_por_pagina
        except Exception:                                  # noqa: BLE001
            logger.debug("Falha ao amostrar texto de %s", file_path, exc_info=True)
            return False   # na dúvida, trata como escaneado

    # -----------------------------------------------------------------
    # Parsing local (só para o que a API não lê)
    # -----------------------------------------------------------------
    def extract_text_from_file(self, file_path):
        nome = os.path.basename(file_path)
        ext = Path(file_path).suffix.lower()
        conteudo = ""

        try:
            if ext in (".xlsx", ".xlsm", ".xltx"):
                from openpyxl import load_workbook
                wb = load_workbook(file_path, read_only=True, data_only=True)
                partes = []
                for aba in wb.worksheets:          # TODAS as abas
                    partes.append(f"\n=== ABA: {aba.title} ===")
                    vazias = 0
                    for linha in aba.iter_rows(values_only=True):
                        celulas = ["" if v is None else str(v).strip() for v in linha]
                        if not any(celulas):
                            vazias += 1
                            if vazias > 25:
                                break
                            continue
                        vazias = 0
                        partes.append(" | ".join(celulas).rstrip(" |"))
                wb.close()
                conteudo = "\n".join(partes)

            elif ext in (".docx", ".dotx"):
                from docx import Document
                d = Document(file_path)
                partes = [p.text for p in d.paragraphs if p.text.strip()]
                for n, tab in enumerate(d.tables, 1):
                    partes.append(f"\n--- tabela {n} ---")
                    for linha in tab.rows:
                        partes.append(" | ".join(c.text.strip() for c in linha.cells))
                conteudo = "\n".join(partes)

            elif ext in (".odt", ".ods"):
                from odf import teletype, text as odftext
                from odf.opendocument import load
                from odf.table import Table, TableCell, TableRow
                d = load(file_path)
                partes = [teletype.extractText(p) for p in d.getElementsByType(odftext.P)]
                for tab in d.getElementsByType(Table):
                    for linha in tab.getElementsByType(TableRow):
                        cel = [teletype.extractText(c) for c in linha.getElementsByType(TableCell)]
                        if any(cel):
                            partes.append(" | ".join(cel))
                conteudo = "\n".join(partes)

            elif ext == ".eml":
                with open(file_path, "rb") as f:
                    msg = email.message_from_binary_file(f, policy=email.policy.default)
                corpo = msg.get_body(preferencelist=("plain", "html"))
                texto = corpo.get_content() if corpo else ""
                if corpo is not None and corpo.get_content_subtype() == "html":
                    texto = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", texto, flags=re.S | re.I)
                    texto = re.sub(r"<[^>]+>", " ", texto)
                anexos = [p.get_filename() or "sem-nome" for p in msg.iter_attachments()]
                conteudo = (f"De: {msg.get('From', '')}\nPara: {msg.get('To', '')}\n"
                            f"Assunto: {msg.get('Subject', '')}\nData: {msg.get('Date', '')}\n"
                            f"Anexos: {', '.join(anexos) or 'nenhum'}\n\n{texto}")

            elif ext == ".xls":
                import xlrd
                aba = xlrd.open_workbook(file_path).sheet_by_index(0)
                partes = [f"\n=== ABA: {aba.name} ==="]
                for r in range(aba.nrows):
                    celulas = []
                    for c in range(aba.ncols):
                        celula = aba.cell(r, c)
                        valor = celula.value
                        if celula.ctype == xlrd.XL_CELL_NUMBER and float(valor).is_integer():
                            valor = int(valor)          # 5330012345678.0 -> 5330012345678
                        celulas.append("" if valor in (None, "") else str(valor).strip())
                    if any(celulas):
                        partes.append(" | ".join(celulas).rstrip(" |"))
                conteudo = "\n".join(partes)

            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    conteudo = f.read()

        except Exception as exc:                           # noqa: BLE001
            logger.exception("Falha ao extrair texto de %s", nome)
            conteudo = f"[erro ao ler {nome}: {type(exc).__name__}: {exc}]"

        conteudo = re.sub(r"\n{3,}", "\n\n", conteudo).strip()
        if len(conteudo) > 300000:
            conteudo = conteudo[:300000] + "\n[conteúdo truncado]"

        return {
            "content": conteudo,
            "mime_type": mimetypes.guess_type(file_path)[0] or "application/octet-stream",
            "filename": nome,
            "size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
        }

    # -----------------------------------------------------------------
    # Chamada à API
    # -----------------------------------------------------------------
    def montar_payload(self, file_path, context, base=None, limite_pi=200):
        blocos, plugins, rota = self.montar_conteudo(file_path)

        # >>> NOVO: lista fechada de PI vinda do modelo de proposta
        lista_pi = ""
        itens_base = (base or {}).get("itens") or []
        if itens_base:
            linhas = [f"{i.get('item')} | {i.get('pi')} | "
                      f"{(i.get('nome_em_portugues') or '')[:80]} | "
                      f"{i.get('qtde') or ''} {i.get('uf') or ''}"
                      for i in itens_base[:limite_pi] if i.get("pi")]
            lista_pi = ("\nLISTA DE PI DO MODELO DE PROPOSTA "
                        "(item | PI | nomenclatura | qtde):\n" + "\n".join(linhas))
            if len(itens_base) > limite_pi:
                lista_pi += f"\n[... {len(itens_base) - limite_pi} itens não listados]"

        # >>> NOVO: avisa o modelo quando o texto veio de OCR local
        origem_ocr = ("\nATENÇÃO: este PDF não tinha camada de texto; ela foi criada "
                      "por OCR local (Tesseract). Trate números com desconfiança, "
                      "conforme a regra 9.") if rota.endswith("+ocr-local") else ""

        blocos.append({"type": "text", "text": (
            f"PROCESSO\n"
            f"- Número: {context.get('numero', 'N/A')}\n"
            f"- Objeto: {context.get('descricao', 'N/A')}\n"
            f"- Valor estimado: {context.get('valor_estimado', 'N/A')}\n"
            f"{lista_pi}\n\n"
            f"ARQUIVO: {os.path.basename(file_path)}{origem_ocr}\n\n"
            f"Extraia os dados no schema abaixo.\n\nSCHEMA:\n{self.ESQUEMA}"
        )})

        payload = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": self.SYSTEM_PROMPT},
                {"role": "user", "content": blocos},
            ],
            "temperature": 0.0,
            "max_tokens": 32000,
            "response_format": {"type": "json_object"},
            # Gemini 3 Flash tem níveis de raciocínio. "low" basta para
            # extração; "minimal" degrada em documento sujo.
            "reasoning": {"effort": "low"},
            "usage": {"include": True},
        }
        if plugins:
            payload["plugins"] = plugins
        return payload, rota

    def process_file(self, file_path, context, base=None):
        # >>> NOVO: proposta em planilha não gasta IA
        if Path(file_path).suffix.lower() in (".xls", ".xlsx", ".xlsm"):
            direto = parse_modelo_proposta(file_path)
            if direto and direto["itens"]:
                return {"status": "success", "data": direto,
                        "uso": {"entrada": 0, "saida": 0, "custo_usd": 0},
                        "rota": "modelo/deterministico"}

        payload, rota = self.montar_payload(file_path, context, base)
        resultado = self._post(payload)
        resultado["rota"] = rota
        return resultado

    def process_with_ai(self, content, context):
        """Compatibilidade: mantém a assinatura antiga (texto puro)."""
        payload = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": self.SYSTEM_PROMPT},
                {"role": "user", "content": (
                    f"PROCESSO\n- Número: {context.get('numero', 'N/A')}\n"
                    f"- Objeto: {context.get('descricao', 'N/A')}\n\n"
                    f"SCHEMA:\n{self.ESQUEMA}\n\nCONTEÚDO:\n{content}"
                )},
            ],
            "temperature": 0.0,
            "max_tokens": 32000,
            "response_format": {"type": "json_object"},
            "reasoning": {"effort": "low"},
            "usage": {"include": True},
        }
        return self._post(payload)

    # -----------------------------------------------------------------
    def _post(self, payload):
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": getattr(settings, "OPENROUTER_SITE_URL", "http://localhost:8000"),
            "X-Title": "Arsenal-Main",
        }

        for tentativa in range(4):
            try:
                r = requests.post(f"{self.base_url}/chat/completions",
                                  headers=headers, json=payload,
                                  timeout=self.timeout,
                                  proxies=self.proxies, verify=self.verify)
            except requests.exceptions.SSLError as exc:
                return {"status": "error", "error":
                        f"TLS rejeitado (CA da rede?): {exc}. "
                        "Configure OPENROUTER_CA_BUNDLE."}
            except (requests.exceptions.ConnectTimeout,
                    requests.exceptions.ProxyError) as exc:
                return {"status": "error", "error":
                        f"sem rota ate openrouter.ai: {exc}. "
                        "Verifique proxy/firewall (HTTPS_PROXY)."}
            except requests.RequestException as exc:
                if tentativa == 3:
                    return {"status": "error", "error": f"rede: {exc}"}
                time.sleep(2 ** tentativa + random.random())
                continue

            if r.status_code == 401:
                return {"status": "error", "error": "chave OpenRouter inválida (401)"}
            if r.status_code == 402:
                return {"status": "error", "error": "sem créditos na OpenRouter (402)"}
            if r.status_code in (408, 429, 500, 502, 503, 504):
                if tentativa == 3:
                    return {"status": "error", "error": f"HTTP {r.status_code} após 4 tentativas"}
                espera = float(r.headers.get("Retry-After") or 2 ** tentativa)
                time.sleep(min(espera, 30) + random.random())
                continue
            if r.status_code >= 400:
                # O corpo da resposta pode conter cabeçalho, prompt e trecho do
                # documento. Vai para o log do servidor, não para a tela.
                logger.error("OpenRouter HTTP %s: %s", r.status_code, r.text[:2000])
                return {"status": "error",
                        "error": f"HTTP {r.status_code} na API de extração"}

            dados = r.json()
            if "error" in dados and not dados.get("choices"):
                return {"status": "error",
                        "error": dados["error"].get("message", "erro da API")}

            escolha = dados["choices"][0]
            texto = escolha["message"].get("content") or ""
            if isinstance(texto, list):
                texto = "".join(b.get("text", "") for b in texto if isinstance(b, dict))

            uso = dados.get("usage", {}) or {}
            try:
                return {
                    "status": "success",
                    "data": extrair_json(texto),
                    "uso": {"entrada": uso.get("prompt_tokens", 0),
                            "saida": uso.get("completion_tokens", 0),
                            "custo_usd": uso.get("cost", 0)},
                    "truncado": escolha.get("finish_reason") == "length",
                    "modelo": dados.get("model", self.model),
                }
            except (ValueError, json.JSONDecodeError) as exc:
                return {"status": "partial", "data": {"conteudo": texto},
                        "error": f"JSON inválido: {exc}"}

        return {"status": "error", "error": "tentativas esgotadas"}

    # -----------------------------------------------------------------
    def process_directory(self, directory_path, context, base=None):
        resultados = []
        for raiz, _dirs, arquivos in os.walk(directory_path):
            for nome in sorted(arquivos):
                caminho = os.path.join(raiz, nome)
                # Link simbólico dentro de um pacote de terceiro é vetor de
                # leitura de arquivo do servidor: o conteúdo apontado seria
                # lido e enviado para a API externa.
                if os.path.islink(caminho) or not os.path.isfile(caminho):
                    logger.warning("Ignorado (não é arquivo regular): %s", caminho)
                    continue
                if os.path.getsize(caminho) > 40 * 1024 * 1024:
                    logger.warning("Ignorado (maior que 40 MB): %s", nome)
                    continue
                if Path(nome).suffix.lower() in (".zip", ".tgz", ".tar", ".gz",
                                                 ".exe", ".dll", ".db"):
                    continue
                resultado = self.process_file(caminho, context, base)
                resultados.append({
                    "filename": nome,
                    "rota": resultado.get("rota", ""),
                    "ai_result": resultado,
                })
        return resultados

    # -----------------------------------------------------------------
    def merge_results(self, results, base=None):
        merged = {"informacoes_gerais": {}, "empresas": [], "itens": [],
                  "perguntas": [], "avisos_gerais": []}
        empresas_vistas = {}
        itens_por_pi = {}
        ordem = []

        def chave(texto):
            return re.sub(r"[^a-z0-9]", "", str(texto or "").lower())

        # --- 1. linha de base define as linhas do mapa -----------------
        for item in (base or {}).get("itens") or []:
            k = chave(item.get("pi"))
            if not k:
                merged["avisos_gerais"].append(
                    f"item {item.get('item')} do modelo sem PI: fora do mapa")
                continue
            itens_por_pi[k] = dict(item, empresas=dict(item.get("empresas") or {}))
            ordem.append(k)
        tem_base = bool(ordem)

        # --- 2. respostas dos fornecedores -----------------------------
        for r in results:
            dados = (r.get("ai_result") or {}).get("data") or {}
            if not isinstance(dados, dict):
                continue
            origem = r.get("filename", "")
            # >>> NOVO: item lido de OCR local não pode passar como certeza
            veio_de_ocr = "ocr-local" in (r.get("rota")
                                          or (r.get("ai_result") or {}).get("rota") or "")

            for k, v in (dados.get("informacoes_gerais") or {}).items():
                if v not in (None, "", []) and not merged["informacoes_gerais"].get(k):
                    merged["informacoes_gerais"][k] = v

            for emp in (dados.get("empresas") or []):
                nome = (emp.get("nome") or "").strip()
                if not nome:
                    continue
                k = chave(emp.get("cnpj")) or chave(nome)
                if k in empresas_vistas:
                    for campo, valor in emp.items():
                        if valor and not empresas_vistas[k].get(campo):
                            empresas_vistas[k][campo] = valor
                    # cotação prevalece sobre declínio/dúvida do mesmo fornecedor
                    if emp.get("tipo_resposta") == "cotacao":
                        empresas_vistas[k]["tipo_resposta"] = "cotacao"
                else:
                    empresas_vistas[k] = dict(emp, nome=nome)

            for item in (dados.get("itens") or []):
                k = chave(item.get("pi")) or chave(item.get("codigo")) or chave(item.get("nsn"))
                precos = {n: v for n, v in (item.get("empresas") or {}).items() if n}

                # >>> NOVO: teto de confiança e aviso quando a origem foi OCR local
                confianca = item.get("confianca", 100)
                avisos_item = list(item.get("avisos") or [])
                if veio_de_ocr:
                    confianca = min(confianca, 70)
                    avisos_item.append(f"valores lidos por OCR local em {origem}: conferir")

                if k in itens_por_pi:
                    alvo = itens_por_pi[k]
                    alvo.setdefault("empresas", {}).update(precos)
                    for campo in ("nsn", "codigo", "nome_em_portugues", "qtde", "uf"):
                        if not alvo.get(campo) and item.get(campo):
                            alvo[campo] = item[campo]
                    alvo["confianca"] = min(alvo.get("confianca", 100), confianca)
                    alvo["avisos"] = (alvo.get("avisos") or []) + avisos_item
                elif tem_base:
                    # PI fora do modelo: não cria linha, só avisa
                    merged["avisos_gerais"].append(
                        f"{origem}: PI '{item.get('pi') or '(vazio)'}' "
                        f"({(item.get('nome_em_portugues') or '')[:60]}) não consta do "
                        f"modelo de proposta — conferir manualmente")
                elif k:
                    itens_por_pi[k] = dict(item, empresas=precos,
                                           confianca=confianca, avisos=avisos_item)
                    ordem.append(k)

            for p in (dados.get("perguntas") or []):
                merged["perguntas"].append(dict(p, arquivo=origem))
            merged["avisos_gerais"].extend(dados.get("avisos_gerais") or [])

        merged["empresas"] = list(empresas_vistas.values())
        merged["itens"] = [itens_por_pi[k] for k in ordem]
        for i, item in enumerate(merged["itens"], 1):
            if not item.get("item"):
                item["item"] = i
        return merged
